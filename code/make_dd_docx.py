"""Build DATA_DICTIONARY.docx: full pipeline + every prompt + the output-table
dictionaries. Prompts are imported live from code/ so they never drift."""
import importlib.util
import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = "/Users/othmanbensouda/Desktop/Data Collection MDL"
CODE = os.path.join(ROOT, "code")
OUT = os.path.join(ROOT, "DATA_DICTIONARY.docx")


def load(modfile, attr):
    spec = importlib.util.spec_from_file_location(modfile, os.path.join(CODE, modfile + ".py"))
    m = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(m)
    return getattr(m, attr)


CLASSIFY_SYS = load("classify_type", "SYS")
REFINE_SYS = load("refine_unclear", "SYS")
GATE_SYS = load("confirm_orders", "SYS")
EXTRACT_SYS = load("extract_orders", "SYSTEM")

# ---------------- pipeline overview ----------------
PIPELINE = [
    ("1", "count_pages.py", "files/ -> page_counts.csv", "free", "Page count per PDF (relpath,pages)."),
    ("2", "classify_type.py", "page_counts.csv -> type_labels.csv", "gpt-5.4-mini",
     "Label each doc Order/Motion/Unclear/Other from filename + page count (conservative)."),
    ("3", "filter_corpus.py", "files/ + type_labels.csv -> filtered_files/", "free",
     "Drop docket dumps, all Other, and Unclear > 50p; hard-link the rest."),
    ("4", "ocr_llamaparse.py", "filtered_files/ -> ocr/<MDL>/<doc>.json", "LlamaParse + Tesseract",
     "OCR to per-page text; 10-page batches; local OCR fallback."),
    ("5", "refine_unclear.py", "type_labels.csv + ocr/ -> type_labels.csv", "gpt-5.4-mini",
     "Rescue real orders sitting in Unclear (regex gate -> LLM -> promote)."),
    ("6", "confirm_orders.py  (GATE)", "ocr/ + type_labels.csv -> order_status.csv", "gpt-5.5",
     "Retrieve only RELEVANT + EXECUTED orders (Test A relevance AND Test B execution)."),
    ("7", "trim_orders.py", "ocr/ -> orders/<MDL>/<doc>.json", "free",
     "Cut each order at its last 'judge' page (drop trailing exhibits)."),
    ("8", "extract_orders.py", "orders/ (retrieve=1) -> order_extractions.{jsonl,xlsx}", "gpt-5.5",
     "Fill the schema per order (+ structured appointees); build the 4-tab workbook."),
]

# ---------------- table dictionaries ----------------
HEADERS = ["Field", "Type", "Values / format", "Source", "Description"]

ORDERS = [
    ("Order_No", "text", "MDL-Docket, e.g. 2263-20", "LLM", "Order identifier; the key Appointments link to."),
    ("MDL_No", "text", "digits", "LLM", "MDL number."),
    ("Docket_No", "integer", "", "LLM", "Docket entry number."),
    ("Date", "text", "ISO yyyy-mm-dd", "LLM", "Date the court issued/entered the order."),
    ("Judge", "text", "initials", "LLM", "Signing judge's initials."),
    ("Judge_Type", "multi-select", "DJ, MJ", "LLM", "District vs magistrate; both if multiple sign."),
    ("Contested", "bool", "", "LLM", "Multiple applicants or objections for a role."),
    ("OU_Create", "integer", "", "LLM", "Count of organizational units created."),
    ("OU_Terminate", "integer", "", "LLM", "Count of organizational units terminated."),
    ("OU_Functions", "bool", "", "LLM", "Order specifies functions for a unit."),
    ("OU_Duties_to_Nonclients", "bool", "", "LLM", "Imposes duties toward non-clients."),
    ("OU_Plaintiff", "bool", "", "LLM", "Affected unit is plaintiff-side."),
    ("OU_Defendant", "bool", "", "LLM", "Affected unit is defense-side."),
    ("Order_Types", "multi-select", "Order_Types vocab", "LLM", "Order-level categories (not generic words)."),
    ("Appointments_Count", "integer", "", "Derived", "Number of linked Appointments rows."),
    ("Applications_Solicited", "bool", "", "LLM", "Court explicitly invited applications, AS STATED IN THIS ORDER (a floor: a separate earlier solicitation order is not captured here)."),
    ("Resolve_Rule_23", "bool", "", "LLM", "Resolves a Rule 23 motion."),
    ("IRPA_Duties_to_Clients", "bool", "", "LLM", "Duties on individually-retained plaintiff attorneys."),
    ("Limit_Nonleader_Practice", "bool", "", "LLM", "Restricts non-lead attorneys' practice."),
    ("Rule_23", "bool", "", "LLM", "Rule 23 cited anywhere."),
    ("MCL", "bool", "", "LLM", "Manual for Complex Litigation cited."),
    ("MDL Type", "multi-select", "free text", "LLM", "JPML classification (often blank)."),
    ("Notes", "long text", "", "LLM", "Freeform notes (e.g. 'Multiple judges')."),
    ("Needs_Motion_Reading", "bool", "TRUE/FALSE/blank", "LLM (Stage 8)",
     "TRUE = order grants an appointment motion without naming appointees -> read the motion. Blank if not an appointment order."),
    ("Needs_Signature_Check", "bool", "TRUE/blank", "Gate (Stage 6)",
     "TRUE = a kept order whose judicial signature was NOT in the OCR text (wet-ink / text / minute order). Genuine order; a human should eyeball the PDF to confirm the signature."),
    ("Source_File", "text", "relpath", "Derived", "Path to the source PDF (provenance)."),
]

APPTS = [
    ("Appointment_ID", "text", "Order_No-N", "Derived", "Synthetic row id (N = position within the order)."),
    ("Order_No", "text", "", "Derived", "Parent order (link key to Orders)."),
    ("Last Name", "text", "", "LLM", "Appointee surname; blank for a firm."),
    ("First Name", "text", "", "LLM", "Appointee given name; blank for a firm."),
    ("Appoint", "bool", "", "LLM", "Being appointed (usual case)."),
    ("Remove", "bool", "", "LLM", "Being removed/terminated from a role."),
    ("Interim", "bool", "", "LLM", "Appointment is interim, AS STATED IN THIS ORDER (a later order making it permanent is not reconciled here)."),
    ("Appointment Types", "multi-select", "Appointment Types vocab", "LLM", "Role(s) given to this appointee."),
    ("Plaintiff/Defendant", "single-select", "Plaintiff, Defendant", "LLM", "Side represented; blank if neutral."),
    ("Appointee Type", "single-select", "Individual, Firm", "LLM", "Person vs law firm."),
    ("Firm", "text", "", "LLM", "Law firm (individual's firm, or the firm itself)."),
    ("First_Last_Calculated", "text", "", "Derived", "'First Last', or the firm name for a firm."),
    ("MDL_No", "text", "", "Derived", "From the parent order."),
    ("MDL Type", "multi-select", "", "Derived", "From the parent order."),
    ("Possible_Duplicate_Appointment", "text", "", "Derived",
     "FLAG only (no row removed): the same appointee+role+side appears in >1 order on the SAME DATE in "
     "this MDL -- a likely duplicate/amended order. ALL rows are kept (gold convention); de-duplication "
     "is a later data-cleaning step. Blank = not a same-date duplicate."),
]

ATTORNEYS = [
    ("Attorney_Identifier", "text", "", "Derived", "Placeholder primary (= auto canonical name)."),
    ("Canonical_Name", "text", "", "Derived -> Human", "Auto-set to 'First Last'/firm; canonicalize manually."),
    ("First_Name", "text", "", "Derived", "From the appointee."),
    ("Last_Name", "text", "", "Derived", "From the appointee."),
    ("Firm", "text", "", "Derived", "From the appointee."),
    ("Gender", "single-select", "M, F, #N/A, 0", "Human", ""),
    ("Race", "single-select", "White; Black or African American; American Indian or Alaska Native; Asian; Native Hawaiian or Other Pacific Islander; Some Other Race", "Human", ""),
    ("Birth_Year", "number", "", "Human", ""),
    ("Undergrad_Grad_Year", "number", "", "Human", ""),
    ("Law_Grad_Year", "number", "", "Human", ""),
    ("Undergrad_School", "text", "", "Human", ""),
    ("Law_School_Name", "text", "", "Human", "Lookup from ABA Law Schools in Airtable."),
    ("Bar_States", "text", "", "Human", ""),
    ("Sources", "long text", "", "Human", ""),
    ("Notes", "long text", "", "Human", ""),
    ("AKA_1 / AKA_2 / AKA_3", "text", "", "Human", "Alternate names used to merge duplicates."),
]

DROPPED = [
    ("Source_File", "text", "relpath", "Derived", "The dropped document's path."),
    ("Relevance", "single-select", "leadership / fees / generic_to_check / irrelevant", "Gate",
     "Test A subject-matter verdict."),
    ("Doc_Kind", "single-select", "executed_order / unsigned_proposed_order / report_recommendation / show_cause / embedded_exhibit_order / motion / ...", "Gate",
     "Test B instrument type."),
    ("Reason", "long text", "", "Gate", "One-sentence rationale for the drop."),
]


def add_table(doc, rows):
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(HEADERS):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    return t


def add_prompt(doc, title, text):
    doc.add_heading(title, level=3)
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line.strip() else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(7.5)


doc = Document()
doc.add_heading("MDL Leadership-Order Pipeline — Data Dictionary & Reference", level=0)
doc.add_paragraph(
    "End-to-end reference for the MDL court-document pipeline: how every stage works, the verbatim "
    "LLM prompts each stage uses, and the field-by-field dictionary for the three Airtable output "
    "tables (plus the Stage-6 audit tab). Source legend for table columns: LLM = produced by the model; "
    "Derived = computed by the export script; Gate = decided by the Stage-6 order gate; Human = left "
    "blank for later research."
)

# ---- Section 1: pipeline ----
doc.add_heading("1. Pipeline overview", level=1)
doc.add_paragraph(
    "Eight stages, each a pure function of the previous stage's durable output, so the whole thing "
    "reproduces from files/. Stages 1, 3, 7 are deterministic/free; 2, 4, 5, 6, 8 are billable and "
    "resumable. Stages 4-8 currently target a 10-MDL sample (seed 281835); extend with --mdls."
)
pt = doc.add_table(rows=1, cols=5)
pt.style = "Light Grid Accent 1"
for i, h in enumerate(["#", "Script", "In -> Out", "Model / cost", "What it does"]):
    c = pt.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
for row in PIPELINE:
    cells = pt.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(val)
        for p in cells[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8.5)

doc.add_paragraph()
doc.add_paragraph(
    "Stage 6 (the gate) is the precision step: it keeps a document only if it passes BOTH a relevance "
    "test (leadership/counsel appointments incl. Rule 23(g) class counsel; attorney's fees / common "
    "benefit incl. settlement orders that award or allocate common-benefit fees; generic CMO/PTO to "
    "check) AND an execution test (a judge actually signed/entered it; signed stipulated and signed "
    "proposed orders count; unsigned proposed orders, magistrate R&Rs, exhibit-embedded orders, "
    "show-cause orders and motions do not). retrieve=1 only when relevance != irrelevant AND the "
    "instrument is an executed order."
)

# ---- Section 2: prompts ----
doc.add_page_break()
doc.add_heading("2. Prompts (verbatim)", level=1)
doc.add_paragraph("The exact system prompts each LLM stage sends, imported live from code/.")
add_prompt(doc, "2.1  Stage 2 — classify_type.py (gpt-5.4-mini)", CLASSIFY_SYS)
add_prompt(doc, "2.2  Stage 5 — refine_unclear.py (gpt-5.4-mini)", REFINE_SYS)
add_prompt(doc, "2.3  Stage 6 — confirm_orders.py — THE GATE (gpt-5.5)", GATE_SYS)
add_prompt(doc, "2.4  Stage 8 — extract_orders.py (gpt-5.5)", EXTRACT_SYS)

# ---- Section 3: vocab ----
doc.add_page_break()
doc.add_heading("3. Controlled vocabularies", level=1)
for line in [
    "Order_Types (15): LeadCounsel, Management, Communications, ClassCounsel, Discovery, Motions, Fees, Expert, Bellwether, Coordination, Settlement, Trial, SettlementAdministration, ProSe, Vetting",
    "Appointment Types (16): the 15 above + LocalCounsel",
    "Judge_Type: DJ (district judge), MJ (magistrate judge)",
    "Plaintiff/Defendant: Plaintiff, Defendant",
    "Appointee Type: Individual, Firm",
]:
    doc.add_paragraph(line, style="List Bullet")

# ---- Section 4: tables ----
doc.add_heading("4. Output tables (order_extractions.xlsx)", level=1)
doc.add_paragraph(
    "Multi-value cells are \"; \"-joined; booleans are real Excel TRUE/FALSE. Airtable-internal fields "
    "(autonumbers, linked records, lookups) are created on import; the export supplies the join keys "
    "(Order_No, MDL_No)."
)
doc.add_heading("4.1  Orders", level=2)
doc.add_paragraph("One row per retrieved order. Primary key: Order_No.")
add_table(doc, ORDERS)
doc.add_heading("4.2  Appointments", level=2)
doc.add_paragraph("One row per appointed/removed person or firm. Links to Orders via Order_No.")
add_table(doc, APPTS)
doc.add_heading("4.3  Attorneys", level=2)
doc.add_paragraph(
    "Deduped roster of appointees (exact first+last or firm only; cross-order variants are merged "
    "manually via Canonical_Name / AKA_*). Demographic fields are blank (external research)."
)
add_table(doc, ATTORNEYS)
doc.add_heading("4.4  Dropped (Stage 6) — audit tab", level=2)
doc.add_paragraph("Every document the gate excluded, with why. Not imported to Airtable; provenance/QA only.")
add_table(doc, DROPPED)

doc.save(OUT)
print("wrote", OUT)
print(f"prompts: classify={len(CLASSIFY_SYS)} refine={len(REFINE_SYS)} gate={len(GATE_SYS)} extract={len(EXTRACT_SYS)} chars")
